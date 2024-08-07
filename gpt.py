from docx import Document

# Create a new Document
doc = Document()

# Title
doc.add_heading('AMD B2B Solutions Overview', 0)

# Section: AMD EPYC™ Processors (Server CPUs)
doc.add_heading('AMD EPYC™ Processors (Server CPUs)', level=1)

# EPYC™ 9004 Series (Genoa)
doc.add_heading('EPYC™ 9004 Series (Genoa)', level=2)
doc.add_paragraph(
    'Key Specifications:\n'
    'Core Count: Up to 96 cores (EPYC 9654)\n'
    'Threads: Up to 192 threads (EPYC 9654)\n'
    'Base Clock: Up to 3.85 GHz (EPYC 9374F)\n'
    'Max Boost Clock: Up to 4.3 GHz (EPYC 9374F)\n'
    'L3 Cache: Up to 384 MB (EPYC 9654)\n'
    'PCIe Lanes: 128 PCIe 5.0 lanes (All models)\n'
    'Memory Support: Up to 6 TB DDR5 (12-channel)\n'
    'TDP: 360W (EPYC 9654 and EPYC 9554)'
)

# EPYC™ 7003 Series
doc.add_heading('EPYC™ 7003 Series', level=2)
doc.add_paragraph(
    'Core Count: Up to 64 cores\n'
    'Threads: Up to 128 threads\n'
    'Base Clock: Up to 3.7 GHz\n'
    'Max Boost Clock: Up to 4.1 GHz\n'
    'L3 Cache: Up to 256 MB\n'
    'PCIe Lanes: 128 PCIe 4.0 lanes\n'
    'Memory Support: Up to 4 TB DDR4 ECC per socket, 8-channel memory'
)

# EPYC™ 7002 Series
doc.add_heading('EPYC™ 7002 Series', level=2)
doc.add_paragraph(
    'Core Count: Up to 64 cores\n'
    'Threads: Up to 128 threads\n'
    'Base Clock: Up to 3.2 GHz\n'
    'Max Boost Clock: Up to 3.4 GHz\n'
    'L3 Cache: Up to 256 MB\n'
    'PCIe Lanes: 128 PCIe 4.0 lanes\n'
    'Memory Support: Up to 4 TB DDR4 ECC per socket, 8-channel memory'
)

# Section: AMD Ryzen™ PRO Processors (Business-Class CPUs)
doc.add_heading('AMD Ryzen™ PRO Processors (Business-Class CPUs)', level=1)

# Ryzen™ PRO 6000 Series
doc.add_heading('Ryzen™ PRO 6000 Series', level=2)
doc.add_paragraph(
    'Core Count: Up to 8 cores\n'
    'Threads: Up to 16 threads\n'
    'Base Clock: Up to 3.3 GHz\n'
    'Max Boost Clock: Up to 4.8 GHz\n'
    'L3 Cache: Up to 16 MB\n'
    'Integrated Graphics: AMD Radeon™ Graphics\n'
    'PCIe Lanes: 24 PCIe 4.0 lanes\n'
    'Memory Support: Up to 64 GB DDR4'
)

# Ryzen™ PRO 5000 Series
doc.add_heading('Ryzen™ PRO 5000 Series', level=2)
doc.add_paragraph(
    'Core Count: Up to 8 cores\n'
    'Threads: Up to 16 threads\n'
    'Base Clock: Up to 3.8 GHz\n'
    'Max Boost Clock: Up to 4.6 GHz\n'
    'L3 Cache: Up to 32 MB\n'
    'PCIe Lanes: 20 PCIe 4.0 lanes\n'
    'Memory Support: Up to 64 GB DDR4'
)

# Ryzen™ PRO 4000 Series
doc.add_heading('Ryzen™ PRO 4000 Series', level=2)
doc.add_paragraph(
    'Core Count: Up to 8 cores\n'
    'Threads: Up to 16 threads\n'
    'Base Clock: Up to 3.0 GHz\n'
    'Max Boost Clock: Up to 4.2 GHz\n'
    'L3 Cache: Up to 8 MB\n'
    'Integrated Graphics: AMD Radeon™ Vega Graphics\n'
    'PCIe Lanes: 24 PCIe 3.0 lanes\n'
    'Memory Support: Up to 64 GB DDR4'
)

# Section: AMD Radeon™ Pro Graphics Cards (Professional GPUs)
doc.add_heading('AMD Radeon™ Pro Graphics Cards (Professional GPUs)', level=1)

# Radeon™ Pro W6000 Series
doc.add_heading('Radeon™ Pro W6000 Series', level=2)
doc.add_paragraph(
    'Stream Processors: Up to 5120\n'
    'Memory: Up to 32 GB GDDR6\n'
    'Memory Bandwidth: Up to 512 GB/s\n'
    'Max Power Consumption: Up to 300 W\n'
    'PCIe Support: PCIe 4.0'
)

# Radeon™ Pro WX Series
doc.add_heading('Radeon™ Pro WX Series', level=2)
doc.add_paragraph(
    'Stream Processors: Up to 2304\n'
    'Memory: Up to 8 GB GDDR5\n'
    'Memory Bandwidth: Up to 256 GB/s\n'
    'Max Power Consumption: Up to 150 W\n'
    'PCIe Support: PCIe 3.0'
)

# Section: AMD Radeon™ Instinct Accelerators (AI & Deep Learning)
doc.add_heading('AMD Radeon™ Instinct Accelerators (AI & Deep Learning)', level=1)

# Radeon™ Instinct MI60
doc.add_heading('Radeon™ Instinct MI60', level=2)
doc.add_paragraph(
    'Compute Units: 64\n'
    'Stream Processors: 4096\n'
    'Memory: 32 GB HBM2\n'
    'Memory Bandwidth: 1 TB/s\n'
    'Max Power Consumption: 300 W\n'
    'PCIe Support: PCIe 4.0'
)

# Radeon™ Instinct MI50
doc.add_heading('Radeon™ Instinct MI50', level=2)
doc.add_paragraph(
    'Compute Units: 60\n'
    'Stream Processors: 3840\n'
    'Memory: 16 GB HBM2\n'
    'Memory Bandwidth: 1 TB/s\n'
    'Max Power Consumption: 300 W\n'
    'PCIe Support: PCIe 4.0'
)

# Section: AMD Ryzen™ Embedded Processors (Embedded Systems)
doc.add_heading('AMD Ryzen™ Embedded Processors (Embedded Systems)', level=1)

# Ryzen™ Embedded V1000 Series
doc.add_heading('Ryzen™ Embedded V1000 Series', level=2)
doc.add_paragraph(
    'Core Count: Up to 4 cores\n'
    'Threads: Up to 8 threads\n'
    'Base Clock: Up to 3.35 GHz\n'
    'Max Boost Clock: Up to 3.8 GHz\n'
    'L3 Cache: Up to 4 MB\n'
    'Integrated Graphics: AMD Radeon™ Vega Graphics\n'
    'TDP: Configurable from 12 W to 54 W\n'
    'Memory Support: Up to 32 GB DDR4'
)

# Ryzen™ Embedded R1000 Series
doc.add_heading('Ryzen™ Embedded R1000 Series', level=2)
doc.add_paragraph(
    'Core Count: Up to 2 cores\n'
    'Threads: Up to 4 threads\n'
    'Base Clock: Up to 3.2 GHz\n'
    'Max Boost Clock: Up to 3.5 GHz\n'
    'L3 Cache: Up to 4 MB\n'
    'Integrated Graphics: AMD Radeon™ Vega Graphics\n'
    'TDP: Configurable from 12 W to 25 W\n'
    'Memory Support: Up to 32 GB DDR4'
)

# Section: Software Solutions
doc.add_heading('Software Solutions', level=1)

# AMD ROCm™ (Radeon Open Compute)
doc.add_heading('AMD ROCm™ (Radeon Open Compute)', level=2)
doc.add_paragraph(
    'Programming Languages: Support for C++, Python, HIP, and more.\n'
    'Libraries: Optimized libraries for deep learning (e.g., MIOpen), linear algebra (e.g., rocBLAS), and other HPC applications.\n'
    'Tools: Developer tools including profilers, debuggers, and compilers.\n'
    'Platform: Open-source platform designed for GPU computing.'
)

# AMD Software: Pro Edition
doc.add_heading('AMD Software: Pro Edition', level=2)
doc.add_paragraph(
    'Features: Professional-grade drivers, advanced rendering settings, and tools tailored for AMD Radeon™ Pro GPUs.\n'
    'Use Cases: Optimized for CAD, media and entertainment, and other professional applications.'
)

# AMD Software: Adrenalin Edition
doc.add_heading('AMD Software: Adrenalin Edition', level=2)
doc.add_paragraph(
    'Features: Gaming-centric drivers and software suite with performance monitoring, game optimization tools, and streaming capabilities.\n'
    'Use Cases: Enhances gaming experience for AMD Radeon™ GPUs.'
)

# Save the document
file_path = "./AMD_B2B_Solutions_Overview.docx"
doc.save(file_path)

